"""
Alternative methods for PDF to DOCX conversion.
This script provides alternative approaches when pdf2docx fails.
"""

import os
import sys
import tempfile
import subprocess
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

def convert_with_pymupdf(pdf_path, docx_path=None):
    """
    Convert PDF to DOCX using PyMuPDF directly.
    This extracts text and images separately and creates a simple DOCX.
    
    Args:
        pdf_path: Path to the PDF file
        docx_path: Path to save the DOCX file (default: same name with .docx extension)
    
    Returns:
        bool: True if conversion was successful, False otherwise
    """
    if not os.path.exists(pdf_path):
        print(f"Error: PDF file not found - {pdf_path}")
        return False
    
    if not docx_path:
        docx_path = os.path.splitext(pdf_path)[0] + "_pymupdf.docx"
    
    print(f"Converting PDF to DOCX using PyMuPDF:")
    print(f"  Input: {pdf_path}")
    print(f"  Output: {docx_path}")
    
    try:
        # Create a new Word document
        doc = Document()
        
        # Open the PDF
        pdf = fitz.open(pdf_path)
        
        # Process each page
        for page_num, page in enumerate(pdf):
            # Add a page header
            doc.add_heading(f"Page {page_num + 1}", level=1)
            
            # Extract text
            text = page.get_text()
            if text.strip():
                doc.add_paragraph(text)
            
            # Extract images
            image_list = page.get_images(full=True)
            
            # Save images to temporary files and add to document
            for img_index, img in enumerate(image_list):
                xref = img[0]
                
                # Extract image
                base_image = pdf.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Save to temp file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                    temp_img.write(image_bytes)
                    temp_img_path = temp_img.name
                
                # Add to document
                try:
                    doc.add_picture(temp_img_path, width=Inches(6))
                    os.unlink(temp_img_path)  # Delete temp file
                except Exception as img_error:
                    print(f"  Warning: Could not add image {img_index}: {str(img_error)}")
        
        # Save the document
        doc.save(docx_path)
        
        # Check if conversion was successful
        if os.path.exists(docx_path) and os.path.getsize(docx_path) > 0:
            print("✓ Conversion successful!")
            return True
        else:
            print("✗ Conversion failed: Output file is empty or not created")
            return False
    
    except Exception as e:
        print(f"✗ Conversion error: {str(e)}")
        return False

def convert_with_external_tools(pdf_path, docx_path=None):
    """
    Try to convert PDF to DOCX using external tools if available.
    
    Args:
        pdf_path: Path to the PDF file
        docx_path: Path to save the DOCX file (default: same name with .docx extension)
    
    Returns:
        bool: True if conversion was successful, False otherwise
    """
    if not os.path.exists(pdf_path):
        print(f"Error: PDF file not found - {pdf_path}")
        return False
    
    if not docx_path:
        docx_path = os.path.splitext(pdf_path)[0] + "_external.docx"
    
    print(f"Attempting conversion with external tools:")
    print(f"  Input: {pdf_path}")
    print(f"  Output: {docx_path}")
    
    # Try LibreOffice if available
    try:
        print("\nTrying LibreOffice...")
        # Check if LibreOffice is installed
        libreoffice_paths = [
            "libreoffice", "soffice",  # Linux/Mac
            r"C:\Program Files\LibreOffice\program\soffice.exe",  # Windows
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"  # Windows 32-bit
        ]
        
        libreoffice_path = None
        for path in libreoffice_paths:
            try:
                subprocess.run([path, "--version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
                libreoffice_path = path
                break
            except (subprocess.SubprocessError, FileNotFoundError):
                continue
        
        if libreoffice_path:
            # Convert using LibreOffice
            subprocess.run(
                [libreoffice_path, "--headless", "--convert-to", "docx", "--outdir", 
                 os.path.dirname(os.path.abspath(docx_path)), pdf_path],
                stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False
            )
            
            # Check if the file was created
            expected_output = os.path.join(
                os.path.dirname(os.path.abspath(docx_path)),
                os.path.splitext(os.path.basename(pdf_path))[0] + ".docx"
            )
            
            if os.path.exists(expected_output):
                # Rename to the desired output path if different
                if expected_output != docx_path:
                    os.rename(expected_output, docx_path)
                print("✓ Conversion with LibreOffice successful!")
                return True
        else:
            print("  LibreOffice not found")
    
    except Exception as e:
        print(f"  LibreOffice conversion failed: {str(e)}")
    
    # If we got here, all external tools failed
    print("✗ External tool conversion failed")
    return False

def try_all_conversion_methods(pdf_path, output_dir=None):
    """
    Try all available conversion methods and report results.
    
    Args:
        pdf_path: Path to the PDF file
        output_dir: Directory to save the converted files (default: same as PDF)
    """
    if not os.path.exists(pdf_path):
        print(f"Error: File not found - {pdf_path}")
        return
    
    if not output_dir:
        output_dir = os.path.dirname(pdf_path) or "."
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    print(f"Trying all conversion methods for: {pdf_path}")
    print(f"Output directory: {output_dir}")
    print("=" * 50)
    
    # Method 1: PyMuPDF direct conversion
    pymupdf_output = os.path.join(output_dir, f"{base_name}_pymupdf.docx")
    pymupdf_success = convert_with_pymupdf(pdf_path, pymupdf_output)
    
    # Method 2: External tools
    external_output = os.path.join(output_dir, f"{base_name}_external.docx")
    external_success = convert_with_external_tools(pdf_path, external_output)
    
    # Summary
    print("\n" + "=" * 50)
    print("SUMMARY OF RESULTS:")
    print("=" * 50)
    
    successful_methods = []
    if pymupdf_success:
        successful_methods.append(("PyMuPDF", pymupdf_output))
    if external_success:
        successful_methods.append(("External Tools", external_output))
    
    if successful_methods:
        print(f"\nSuccessful conversions: {len(successful_methods)}")
        for method, path in successful_methods:
            print(f"- {method}: {path}")
    else:
        print("\nNo successful conversions. All methods failed.")
        print("The PDF file might be incompatible or protected.")
        print("Consider using an online conversion service.")

if __name__ == "__main__":
    # Check if PDF file is provided as argument
    if len(sys.argv) < 2:
        print("Usage: python alternative_conversion.py <pdf_file> [output_directory]")
        sys.exit(1)
    
    pdf_file = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    try_all_conversion_methods(pdf_file, output_dir) 